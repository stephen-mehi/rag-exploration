from docx import Document as DocxDocument  # <-- Rename this to avoid collision
import os

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
from langchain.docstore.document import Document as LangchainDocument

from langchain.llms import Ollama
from langchain.chains import RetrievalQA

import gradio as gr


def load_documents_from_folder(folder_path):
    docs = []
    for filename in os.listdir(folder_path):
        if filename.endswith(".docx") and not filename.startswith("~$"):
            path = os.path.join(folder_path, filename)
            doc = DocxDocument(path)
            text = "\n".join([p.text for p in doc.paragraphs])
            docs.append(text)
    
    return docs


# Load and chunk
text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=100)
documents = [LangchainDocument(page_content=doc) for doc in load_documents_from_folder("C:/Users/SMehi/Work/onboarding/AgileTraining")]
chunks = text_splitter.split_documents(documents)

# Embed and store
embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
vectorstore = FAISS.from_documents(chunks, embeddings)


# build langchain
llm = Ollama(model="mistral")  # Or "llama2", "gemma", etc.

qa = RetrievalQA.from_chain_type(
    llm=llm,
    retriever=vectorstore.as_retriever(),
    return_source_documents=True
)



# Gradio UI
def answer_query(query):
    response = qa({"query": query})
    result = response["result"]
    sources = [doc.metadata["source"] for doc in response["source_documents"]]
    sources_text = "\n".join(set(sources))
    return f"Answer:\n{result}\n\n Sources:\n{sources_text}"

gr.Interface(fn=answer_query, inputs="text", outputs="text", title="Local RAG Q&A").launch()

